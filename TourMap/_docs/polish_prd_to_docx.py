import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE_MD = Path(r"D:\Code\ProjectCSharp\TourMap\_docs\AudioTourApp_PRD_v2.md")
OUTPUT_MD = Path(r"D:\Code\ProjectCSharp\TourMap\_docs\AudioTourApp_PRD_v2_polished.md")
OUTPUT_DOCX = Path(r"D:\Code\ProjectCSharp\TourMap\_docs\AudioTourApp_PRD_v2_polished.docx")


def normalize_text(text: str) -> str:
    # Remove emoji and decorative symbols to keep an academic PRD tone.
    text = re.sub(r"[\U0001F300-\U0001FAFF]", "", text)
    text = text.replace("✅", "").replace("⚙️", "").replace("🔧", "")

    replacements = {
        "Skip": "Bỏ qua",
        "bấm": "chọn",
        "Bấm": "Chọn",
        "Không làm gì": "Tiếp tục theo dõi",
        "trái tim của ứng dụng": "thành phần lõi của ứng dụng",
        "chống spam": "tránh kích hoạt lặp",
        "MỚI": "CẬP NHẬT",
        "preview": "xem trước",
        "Hot reload": "hỗ trợ gỡ lỗi nhanh",
        "Polish": "Hoàn thiện",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)

    # Contextual replacements to avoid changing product names and URLs.
    text = re.sub(r"\bMobile App\b", "Ứng dụng di động", text)
    text = re.sub(r"\bAdmin App\b", "Ứng dụng quản trị", text)
    text = re.sub(r"\bmở App\b", "mở ứng dụng", text)
    text = re.sub(r"\bkhởi động App\b", "khởi động ứng dụng", text, flags=re.IGNORECASE)

    # Standardize some common terms.
    text = text.replace("offline-first", "offline first")
    text = text.replace("Offline-First", "Offline First")
    text = text.replace("REST API", "RESTful API")
    text = text.replace("real-time", "thời gian thực")
    text = text.replace("geofence", "geofence")

    # Tidy repeated spaces introduced by replacements.
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text


def apply_markdown_polish(md: str) -> str:
    md = normalize_text(md)

    # Tone refinements for key recurring patterns.
    md = md.replace(
        "Ứng dụng hướng tới việc số hóa trải nghiệm tham quan, thay thế tour guide truyền thống bằng công nghệ định vị và âm thanh hiện đại, phù hợp với nhu cầu du lịch tự do và các tuyến di sản văn hóa đô thị.",
        "Sản phẩm hướng đến số hóa trải nghiệm tham quan thông qua công nghệ định vị và âm thanh số, qua đó hỗ trợ mô hình du lịch tự khám phá tại các tuyến di sản văn hóa đô thị.",
    )
    md = md.replace(
        "Mọi dữ liệu POI, audio và bản đồ đều được cache cục bộ trên SQLite & file system, đảm bảo ứng dụng vận hành 100% khi mất kết nối.",
        "Dữ liệu POI, nội dung âm thanh và dữ liệu hỗ trợ hiển thị được lưu đệm cục bộ trên SQLite và hệ thống tệp, bảo đảm các chức năng cốt lõi vẫn hoạt động ổn định khi mất kết nối mạng.",
    )
    md = md.replace(
        "Hệ thống CMS là ứng dụng web dành cho người quản trị, cho phép quản lý toàn bộ nội dung, theo dõi analytics và cấu hình hệ thống mà không cần can thiệp vào code.",
        "Hệ thống CMS là nền tảng web phục vụ quản trị nội dung, theo dõi phân tích dữ liệu và cấu hình vận hành mà không yêu cầu chỉnh sửa mã nguồn trực tiếp.",
    )

    return md


def set_default_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph("AUDIO TOUR APP")
    p.style = doc.styles["Title"]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p2 = doc.add_paragraph("Product Requirements Document (PRD) - Phiên bản đã biên tập")
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p3 = doc.add_paragraph("Tháng 3, 2026")
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()


def parse_table(lines, i):
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append(lines[i].strip())
        i += 1
    return rows, i


def write_table(doc: Document, rows):
    if len(rows) < 2:
        return
    header = [c.strip() for c in rows[0].strip("|").split("|")]
    data_rows = rows[2:] if len(rows) >= 2 else []

    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, col in enumerate(header):
        hdr_cells[idx].text = col

    for row in data_rows:
        cols = [c.strip() for c in row.strip("|").split("|")]
        tr = table.add_row().cells
        for j in range(min(len(cols), len(header))):
            tr[j].text = cols[j]

    doc.add_paragraph("")


def markdown_to_docx(md_text: str, output_path: Path) -> None:
    doc = Document()
    set_default_style(doc)
    add_title_block(doc)

    lines = md_text.splitlines()
    i = 0
    in_code = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            if in_code:
                p = doc.add_paragraph("Khối mã / sơ đồ:")
                p.runs[0].bold = True
            else:
                doc.add_paragraph("")
            i += 1
            continue

        if in_code:
            p = doc.add_paragraph(line)
            p.style = doc.styles["No Spacing"]
            i += 1
            continue

        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            write_table(doc, rows)
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif stripped == "---":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)
        i += 1

    doc.save(output_path)


def main():
    if not SOURCE_MD.exists():
        raise FileNotFoundError(f"Không tìm thấy file nguồn: {SOURCE_MD}")

    md = SOURCE_MD.read_text(encoding="utf-8")
    polished_md = apply_markdown_polish(md)
    OUTPUT_MD.write_text(polished_md, encoding="utf-8")
    markdown_to_docx(polished_md, OUTPUT_DOCX)

    print(f"Generated: {OUTPUT_MD}")
    print(f"Generated: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
