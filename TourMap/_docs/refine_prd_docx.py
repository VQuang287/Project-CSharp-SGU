import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


INPUT_DOCX = Path(r"D:\Code\ProjectCSharp\TourMap\_docs\AudioTourApp_PRD_v2.docx")
OUTPUT_DOCX = Path(r"D:\Code\ProjectCSharp\TourMap\_docs\AudioTourApp_PRD_v2_refined.docx")


def normalize(s: str) -> str:
    # Remove accents for robust keyword matching.
    mapping = str.maketrans(
        "àáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđ",
        "aaaaaaaaaaaaaaaaaeeeeeeeeeeeiiiiiooooooooooooooooouuuuuuuuuuuyyyyyd",
    )
    return s.lower().translate(mapping).strip()


def remove_unneeded_metadata_rows(doc: Document) -> None:
    keys_to_drop = {"phien ban", "trang thai"}
    for table in doc.tables:
        rows_to_remove = []
        for row in table.rows:
            key = normalize(row.cells[0].text if row.cells else "")
            if key in keys_to_drop:
                rows_to_remove.append(row)
        for row in rows_to_remove:
            table._tbl.remove(row._tr)


def apply_heading_styles(doc: Document) -> None:
    h1 = re.compile(r"^\d+\.\s+.+")
    h2 = re.compile(r"^\d+\.\d+\s+.+")
    h3 = re.compile(r"^\d+\.\d+\.\d+\s+.+")

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        # Apply most specific first.
        if h3.match(text):
            p.style = doc.styles["Heading 3"]
        elif h2.match(text):
            p.style = doc.styles["Heading 2"]
        elif h1.match(text):
            p.style = doc.styles["Heading 1"]


def clear_toc_block_and_insert_auto_toc(doc: Document) -> None:
    toc_heading_idx = None
    for i, p in enumerate(doc.paragraphs):
        if normalize(p.text) in {"muc luc", "mc lc"}:
            toc_heading_idx = i
            break
    if toc_heading_idx is None:
        return

    start = toc_heading_idx + 1
    end = start
    for j in range(start, len(doc.paragraphs)):
        t = doc.paragraphs[j].text.strip()
        if re.match(r"^\d+\.\s+.+", t):
            end = j
            break
    else:
        end = len(doc.paragraphs)

    # Remove existing static TOC content (danh sách mục lục cũ).
    for j in range(end - 1, start - 1, -1):
        doc.paragraphs[j]._element.getparent().remove(doc.paragraphs[j]._element)

    # Không chèn TOC field nữa để tương thích tốt với Google Docs.
    # Chỉ để trống sau heading "MỤC LỤC", người dùng sẽ Insert → Table of contents bên phía Google Docs.


def main() -> None:
    if not INPUT_DOCX.exists():
        raise FileNotFoundError(f"Input file not found: {INPUT_DOCX}")

    doc = Document(INPUT_DOCX)
    remove_unneeded_metadata_rows(doc)
    apply_heading_styles(doc)
    clear_toc_block_and_insert_auto_toc(doc)
    doc.save(OUTPUT_DOCX)
    print(f"Generated: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
